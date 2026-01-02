#!/usr/bin/env python3
"""
Hotel Nametag Generator
-----------------------
Extracts guest data from hotel arrival PDF and generates printable nametags in Word format.

Usage: Run the script, select PDF via file dialog, save output .docx file.
"""

import re
import os
import sys
from dataclasses import dataclass

# Conditional tkinter import (for environments without display)
try:
    import tkinter as tk
    from tkinter import filedialog, messagebox
    TKINTER_AVAILABLE = True
except ImportError:
    TKINTER_AVAILABLE = False
from typing import List, Tuple
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# QR code generation
import qrcode
from io import BytesIO


# ============================================================================
# DATA STRUCTURES
# ============================================================================

@dataclass
class GuestRecord:
    """Represents a single guest entry for nametag generation."""
    id: str
    last_name: str
    room_number: str
    number_of_guests: int
    arrival_day: str
    departure_day: str
    matchcode: str = ""  # Booking reference (e.g., "XXD/2521/1")


# ============================================================================
# PDF PARSER
# ============================================================================

class PDFParser:
    """Extracts and parses guest data from hotel arrival PDF."""
    
    # Regex patterns
    # Matches room patterns like: 001(1), -1(2), 502-C(1), 210 T(2)
    # Must be preceded by whitespace or start of string to avoid matching matchcode parts like /1
    ROOM_PATTERN = re.compile(r'(?:^|\s)(-?\d+(?:-[A-Z])?|\d+\s[A-Z])\((\d+)\)')
    DATE_PATTERN = re.compile(r'(\d{2})\.\d{2}\.\d{2}')  # Matches: DD.MM.YY
    
    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.raw_lines: List[str] = []
        self.guest_records: List[GuestRecord] = []
        
    def extract_text(self) -> None:
        """Extract text lines from PDF."""
        with pdfplumber.open(self.pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    self.raw_lines.extend(text.split('\n'))
    
    def parse_matchcode(self, matchcode: str) -> str:
        """Extract last name from matchcode (everything before first underscore)."""
        if '_' in matchcode:
            return matchcode.split('_')[0].upper()
        # Handle cases without underscore
        return matchcode.split('/')[0].upper() if '/' in matchcode else matchcode.upper()

    def extract_booking_code(self, matchcode: str) -> str:
        """
        Extract booking reference code from matchcode string.
        E.g., 'AHMAD_WASEEM_XXD/2521/1' -> 'XXD/2521/1'
        """
        if '/' not in matchcode:
            return ""

        # Find the position of the first '/'
        slash_pos = matchcode.find('/')

        # Look backwards from the slash to find the last underscore before it
        last_underscore = matchcode.rfind('_', 0, slash_pos)

        if last_underscore != -1:
            # Return everything after the last underscore before the slash
            return matchcode[last_underscore + 1:].strip()
        else:
            # No underscore before slash, return from start
            return matchcode.strip()
    
    def parse_rooms(self, room_str: str) -> List[Tuple[str, int]]:
        """
        Parse room string into list of (room_number, guest_count) tuples.
        Handles: "102(4)", "-1(2)", "502-C(1)", "210 T(2)"
        """
        rooms = []
        matches = self.ROOM_PATTERN.findall(room_str)
        for room_id, count in matches:
            room_id = room_id.strip()
            rooms.append((room_id, int(count)))
        return rooms
    
    def parse_date(self, date_str: str) -> str:
        """Extract DD from DD.MM.YY format."""
        match = self.DATE_PATTERN.search(date_str)
        return match.group(1) if match else ""
    
    def is_continuation_line(self, line: str) -> bool:
        """Check if line is a room continuation (no matchcode, just room data)."""
        line = line.strip()
        if not line:
            return False
        
        parts = line.split()
        
        # Check if line has rooms
        has_rooms = bool(self.ROOM_PATTERN.search(line))
        
        # Check if line has dates (main data lines have dates)
        dates = [p for p in parts if self.DATE_PATTERN.match(p)]
        
        # Continuation line has rooms but no dates
        if has_rooms and len(dates) == 0:
            return True
        
        # Single O or X markers
        if line in ('O', 'X'):
            return True
        
        return False
    
    def parse_data_line(self, line: str) -> dict:
        """
        Parse a main data line into components.
        Returns dict with: matchcode, rooms, arrival, departure
        """
        parts = line.split()
        if len(parts) < 5:
            return None
            
        result = {
            'matchcode': '',
            'rooms': [],
            'arrival': '',
            'departure': ''
        }
        
        # First part is matchcode
        result['matchcode'] = parts[0]
        
        # Parse rooms directly from the entire line
        result['rooms'] = self.parse_rooms(line)
        
        # Parse dates (first date = arrival, second = departure)
        dates_found = [p for p in parts if self.DATE_PATTERN.match(p)]
        if len(dates_found) >= 2:
            result['arrival'] = self.parse_date(dates_found[0])
            result['departure'] = self.parse_date(dates_found[1])
        
        return result
    
    def process(self) -> List[GuestRecord]:
        """Main processing pipeline. Returns list of GuestRecord objects."""
        self.extract_text()
        
        current_entry = None
        guest_id = 0
        
        for line in self.raw_lines:
            line = line.strip()
            
            # Skip header lines, empty lines, page markers
            if not line or 'Matchcode' in line or 'Arrivals' in line or 'Seite' in line:
                continue
            if line.startswith('EHC/') or line.startswith('Room Arrival'):
                continue
            # Skip summary lines (just numbers at end)
            if line.isdigit():
                continue
                
            # Check if this is a continuation line
            if self.is_continuation_line(line) and current_entry:
                # Add rooms from continuation
                additional_rooms = self.parse_rooms(line)
                current_entry['rooms'].extend(additional_rooms)
                continue
            
            # Try to parse as main data line
            parsed = self.parse_data_line(line)
            if parsed and parsed['matchcode'] and parsed['rooms']:
                # Save previous entry if exists
                if current_entry and current_entry['rooms']:
                    self._create_guest_records(current_entry, guest_id)
                    guest_id += sum(count for _, count in current_entry['rooms'])
                
                current_entry = parsed
        
        # Don't forget the last entry
        if current_entry and current_entry['rooms']:
            self._create_guest_records(current_entry, guest_id)
        
        return self.guest_records
    
    def _create_guest_records(self, entry: dict, start_id: int) -> None:
        """Create GuestRecord objects for each guest in an entry."""
        last_name = self.parse_matchcode(entry['matchcode'])
        booking_code = self.extract_booking_code(entry['matchcode'])
        arrival = entry['arrival']
        departure = entry['departure']

        current_id = start_id
        for room_number, guest_count in entry['rooms']:
            for _ in range(guest_count):
                record = GuestRecord(
                    id=f"{current_id:03d}",
                    last_name=last_name,
                    room_number=room_number,
                    number_of_guests=guest_count,
                    arrival_day=arrival,
                    departure_day=departure,
                    matchcode=booking_code
                )
                self.guest_records.append(record)
                current_id += 1


# ============================================================================
# WORD DOCUMENT GENERATOR
# ============================================================================

class NametageGenerator:
    """Generates Word document with nametags from guest records."""
    
    # Layout constants - A4: 21cm x 29.7cm
    TAGS_PER_ROW = 2
    ROWS_PER_PAGE = 6
    TAGS_PER_PAGE = 12
    
    # Cell dimensions - A4: 21cm x 29.7cm
    CELL_WIDTH_CM = 10.5  # Width per cell (21cm / 2)
    ROW_HEIGHT_CM = 4.80  # Height per row
    
    # WiFi configuration
    WIFI_SSID = "Czech-Inn-Public-NEW"
    WIFI_PASSWORD = "iczechedinn"
    
    def __init__(self, guests: List[GuestRecord]):
        self.guests = guests
        self.doc = Document()
        self._setup_document()
        self._qr_cache = None  # Cache QR code bytes
    
    def _generate_wifi_qr(self) -> BytesIO:
        """Generate QR code for WiFi connection."""
        if self._qr_cache:
            self._qr_cache.seek(0)
            return self._qr_cache
        
        # WiFi QR code format
        wifi_string = f"WIFI:T:WPA;S:{self.WIFI_SSID};P:{self.WIFI_PASSWORD};;"
        
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=1,
        )
        qr.add_data(wifi_string)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)
        self._qr_cache = buffer
        return buffer
    
    def _setup_document(self) -> None:
        """Configure document margins and page size for A4."""
        section = self.doc.sections[0]
        section.page_width = Cm(21)  # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0)
    
    def _set_cell_border(self, cell, border_type="dashed", color="000000", size=4):
        """Set dashed border on a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_type)
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    def _add_nametag_content(self, cell, guest: GuestRecord) -> None:
        """Fill a table cell with nametag content for a guest."""
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        
        def set_spacing(para, before=0, after=0):
            """Set compact paragraph spacing."""
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), str(int(before * 20)))
            spacing.set(qn('w:after'), str(int(after * 20)))
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
        
        def clear_borders(tbl_cell):
            """Remove borders from a table cell."""
            tc = tbl_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)
        
        def set_cell_margins(tbl_cell, left=0, right=0, top=0, bottom=0):
            """Set cell margins in twips."""
            tc = tbl_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side, val in [('left', left), ('right', right), ('top', top), ('bottom', bottom)]:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), str(val))
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
        
        # Delete all existing paragraphs and start fresh
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)
        
        # Set cell vertical alignment to top
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # ============================================================
        # ROW 1: CZECHINN + room number (centered together)
        # ============================================================
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p1, 3, 2)
        
        # CZECHINN (Calibri, 18pt)
        r_hotel = p1.add_run("CZECHINN")
        r_hotel.font.name = "Calibri"
        r_hotel._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r_hotel.font.size = Pt(18)

        # Space + Room number placeholder (blank for manual filling)
        p1.add_run("  ")  # Small space between
        r_room = p1.add_run("# ______")
        r_room.bold = True
        r_room.font.size = Pt(16)
        
        # ============================================================
        # ROW 2: Reservation name (centered)
        # ============================================================
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, 2, 2)
        
        r_label = p2.add_run("Reservation name: ")
        r_label.font.size = Pt(10)
        
        r_value = p2.add_run(guest.last_name)
        r_value.font.size = Pt(12)
        r_value.bold = True
        
        # ============================================================
        # ROW 3: QR (left) | Info (right) - nested table
        # ============================================================
        tbl = cell.add_table(rows=1, cols=2)
        tbl.autofit = False
        tbl.allow_autofit = False
        
        qr_cell = tbl.cell(0, 0)
        info_cell = tbl.cell(0, 1)
        qr_cell.width = Cm(3.2)
        info_cell.width = Cm(6.3)
        clear_borders(qr_cell)
        clear_borders(info_cell)
        set_cell_margins(qr_cell, 0, 0, 0, 0)
        set_cell_margins(info_cell, 100, 0, 0, 0)
        
        # QR code (larger)
        qr_para = qr_cell.paragraphs[0]
        qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(qr_para, 0, 0)
        qr_buffer = self._generate_wifi_qr()
        qr_buffer.seek(0)
        qr_para.add_run().add_picture(qr_buffer, height=Cm(2.2))
        
        # Info - 3 lines with 1.5 line spacing
        def set_spacing_15(para, before=0, after=0):
            """Set paragraph spacing with 1.5 line spacing."""
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), str(int(before * 20)))
            spacing.set(qn('w:after'), str(int(after * 20)))
            spacing.set(qn('w:line'), '360')  # 1.5 lines = 360 twips
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
        
        # Line 1: From/To (12pt) with bold dates
        i1 = info_cell.paragraphs[0]
        set_spacing_15(i1, 0, 0)
        r_from = i1.add_run("From: ")
        r_from.font.size = Pt(12)
        r_from_val = i1.add_run(guest.arrival_day)
        r_from_val.font.size = Pt(12)
        r_from_val.bold = True
        r_to = i1.add_run("  To: ")
        r_to.font.size = Pt(12)
        r_to_val = i1.add_run(guest.departure_day)
        r_to_val.font.size = Pt(12)
        r_to_val.bold = True
        
        # Line 2: Check-out time
        i2 = info_cell.add_paragraph()
        set_spacing_15(i2, 0, 0)
        i2.add_run("Check-out time: 12:00").font.size = Pt(9)
        
        # Line 3: WiFi code
        i3 = info_cell.add_paragraph()
        set_spacing_15(i3, 0, 0)
        i3.add_run(f"wifi code: {self.WIFI_PASSWORD}").font.size = Pt(9)
        
        # ============================================================
        # ROW 4: ID value only - remove any trailing empty paragraphs first
        # ============================================================
        # The nested table adds an empty paragraph - find and remove it
        for para in cell.paragraphs:
            if not para.text.strip() and para != cell.paragraphs[0]:
                para._element.getparent().remove(para._element)
        
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p4, 3, 0)
        p4.add_run(guest.id).font.size = Pt(8)
    
    def _add_empty_nametag(self, cell) -> None:
        """Fill a table cell with an empty nametag template for manual filling."""
        from docx.shared import Pt, Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        
        def set_spacing(para, before=0, after=0):
            """Set compact paragraph spacing."""
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), str(int(before * 20)))
            spacing.set(qn('w:after'), str(int(after * 20)))
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
        
        def set_spacing_15(para, before=0, after=0):
            """Set paragraph spacing with 1.5 line spacing."""
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), str(int(before * 20)))
            spacing.set(qn('w:after'), str(int(after * 20)))
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
        
        def clear_borders(tbl_cell):
            """Remove borders from a table cell."""
            tc = tbl_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)
        
        def set_cell_margins(tbl_cell, left=0, right=0, top=0, bottom=0):
            """Set cell margins in twips."""
            tc = tbl_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for side, val in [('left', left), ('right', right), ('top', top), ('bottom', bottom)]:
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), str(val))
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
        
        # Delete all existing paragraphs
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)
        
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # ROW 1: CZECHINN + blank room number
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p1, 3, 2)
        
        r_hotel = p1.add_run("CZECHINN")
        r_hotel.font.name = "Calibri"
        r_hotel._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r_hotel.font.size = Pt(18)
        
        p1.add_run("  ")
        r_room = p1.add_run("# ______")
        r_room.bold = True
        r_room.font.size = Pt(16)
        
        # ROW 2: Reservation name with blank
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, 2, 2)
        
        r_label = p2.add_run("Reservation name: ")
        r_label.font.size = Pt(10)
        
        r_value = p2.add_run("_________________")
        r_value.font.size = Pt(12)
        r_value.bold = True
        
        # ROW 3: QR + Info with blanks
        tbl = cell.add_table(rows=1, cols=2)
        tbl.autofit = False
        tbl.allow_autofit = False
        
        qr_cell = tbl.cell(0, 0)
        info_cell = tbl.cell(0, 1)
        qr_cell.width = Cm(3.2)
        info_cell.width = Cm(6.3)
        clear_borders(qr_cell)
        clear_borders(info_cell)
        set_cell_margins(qr_cell, 0, 0, 0, 0)
        set_cell_margins(info_cell, 100, 0, 0, 0)
        
        # QR code
        qr_para = qr_cell.paragraphs[0]
        qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(qr_para, 0, 0)
        qr_buffer = self._generate_wifi_qr()
        qr_buffer.seek(0)
        qr_para.add_run().add_picture(qr_buffer, height=Cm(2.2))
        
        # Info with blanks for dates
        i1 = info_cell.paragraphs[0]
        set_spacing_15(i1, 0, 0)
        r_from = i1.add_run("From: ")
        r_from.font.size = Pt(12)
        r_from_val = i1.add_run("____")
        r_from_val.font.size = Pt(12)
        r_from_val.bold = True
        r_to = i1.add_run("  To: ")
        r_to.font.size = Pt(12)
        r_to_val = i1.add_run("____")
        r_to_val.font.size = Pt(12)
        r_to_val.bold = True
        
        i2 = info_cell.add_paragraph()
        set_spacing_15(i2, 0, 0)
        i2.add_run("Check-out time: 12:00").font.size = Pt(9)
        
        i3 = info_cell.add_paragraph()
        set_spacing_15(i3, 0, 0)
        i3.add_run(f"wifi code: {self.WIFI_PASSWORD}").font.size = Pt(9)
        
        # ROW 4: No ID for empty nametags
        for para in cell.paragraphs:
            if not para.text.strip() and para != cell.paragraphs[0]:
                para._element.getparent().remove(para._element)
    
    def _create_page_table(self) -> 'Table':
        """Create a 6x2 table for one page of nametags."""
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        
        table = self.doc.add_table(rows=self.ROWS_PER_PAGE, cols=self.TAGS_PER_ROW)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for col in table.columns:
            for cell in col.cells:
                cell.width = Cm(self.CELL_WIDTH_CM)
        
        # Set row heights (EXACTLY - prevents row expansion)
        for row in table.rows:
            row.height = Cm(self.ROW_HEIGHT_CM)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                self._set_cell_border(cell)
        
        return table
    
    def generate(self, output_path: str) -> int:
        """
        Generate the Word document with all nametags.
        Returns the number of pages generated.
        """
        total_guests = len(self.guests)
        total_pages = (total_guests + self.TAGS_PER_PAGE - 1) // self.TAGS_PER_PAGE

        guest_index = 0

        for page_num in range(total_pages):
            # Add page break before new pages (not first page)
            if page_num > 0:
                self.doc.add_page_break()

            # Create table for this page
            table = self._create_page_table()

            # Fill cells column by column (top-to-bottom in left column, then right column)
            # Left column: 000-005, Right column: 006-011
            for col_idx in range(self.TAGS_PER_ROW):
                for row_idx in range(self.ROWS_PER_PAGE):
                    cell = table.cell(row_idx, col_idx)
                    if guest_index < total_guests:
                        self._add_nametag_content(cell, self.guests[guest_index])
                        guest_index += 1
                    else:
                        # Fill remaining cells with empty nametags
                        self._add_empty_nametag(cell)

        # Save document
        self.doc.save(output_path)
        return total_pages


# ============================================================================
# GUI APPLICATION
# ============================================================================

if TKINTER_AVAILABLE:
    class Application:
        """Main GUI application for the nametag generator."""
        
        def __init__(self):
            self.root = tk.Tk()
            self.root.withdraw()  # Hide main window
        
        def select_pdf(self) -> str:
            """Open file dialog to select input PDF."""
            filepath = filedialog.askopenfilename(
                title="Select Hotel Arrivals PDF",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )
            return filepath
        
        def select_output(self) -> str:
            """Open file dialog to select output location."""
            filepath = filedialog.asksaveasfilename(
                title="Save Nametags Document",
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
                initialfile="nametags_output.docx"
            )
            return filepath
        
        def show_info(self, title: str, message: str) -> None:
            """Show information dialog."""
            messagebox.showinfo(title, message)
        
        def show_error(self, title: str, message: str) -> None:
            """Show error dialog."""
            messagebox.showerror(title, message)
        
        def run(self) -> None:
            """Main application flow."""
            # Step 1: Select PDF
            pdf_path = self.select_pdf()
            if not pdf_path:
                self.show_info("Cancelled", "No PDF file selected. Exiting.")
                return
            
            # Step 2: Parse PDF
            try:
                parser = PDFParser(pdf_path)
                guests = parser.process()
                
                if not guests:
                    self.show_error("Error", "No guest records found in the PDF.")
                    return
                    
            except Exception as e:
                self.show_error("PDF Error", f"Failed to parse PDF:\n{str(e)}")
                return
            
            # Step 3: Select output location
            output_path = self.select_output()
            if not output_path:
                self.show_info("Cancelled", "No output location selected. Exiting.")
                return
            
            # Step 4: Generate document
            try:
                generator = NametageGenerator(guests)
                pages = generator.generate(output_path)
                
                self.show_info(
                    "Success",
                    f"Nametags generated successfully!\n\n"
                    f"Total guests: {len(guests)}\n"
                    f"Total pages: {pages}\n"
                    f"Output: {output_path}"
                )
                
            except Exception as e:
                self.show_error("Generation Error", f"Failed to generate document:\n{str(e)}")
                return
            
            self.root.destroy()


# ============================================================================
# CLI MODE (fallback when no GUI available)
# ============================================================================

def run_cli(pdf_path: str, output_path: str) -> None:
    """Command-line interface for batch processing."""
    print(f"Processing: {pdf_path}")
    
    # Parse PDF
    parser = PDFParser(pdf_path)
    guests = parser.process()
    
    if not guests:
        print("ERROR: No guest records found in the PDF.")
        return
    
    print(f"Found {len(guests)} guest records")
    
    # Generate document
    generator = NametageGenerator(guests)
    pages = generator.generate(output_path)
    
    print(f"Generated {pages} pages -> {output_path}")


# ============================================================================
# ENTRY POINT
# ============================================================================

def main():
    """Main entry point."""
    # Check for command line arguments (CLI mode)
    if len(sys.argv) >= 3:
        pdf_path = sys.argv[1]
        output_path = sys.argv[2]
        run_cli(pdf_path, output_path)
        return
    
    # GUI mode
    if TKINTER_AVAILABLE:
        app = Application()
        app.run()
    else:
        print("Usage: python hotel_nametag_generator.py <input.pdf> <output.docx>")
        print("\nGUI mode requires tkinter. Use command line arguments instead.")


if __name__ == "__main__":
    main()
